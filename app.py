"""Sci-RAG application entrypoint.

Importing this module is intentionally side-effect free.  Models, the OpenAI
client, ChromaDB, and Gradio are created only by :func:`create_runtime` or
when the UI is launched from ``main``.  The parsing and table logic lives in
``sci_rag_core.py`` so it can be tested offline.
"""

from __future__ import annotations

from dataclasses import dataclass
import hashlib
import html
import os
from pathlib import Path
import random
import re
from typing import Any

from sci_rag_core import (
    Chunk,
    build_evidence_ledger,
    extract_spatial_figure_chunks,
    formula_evidence_indices,
    find_table_cell_in_chunks,
    figure_reference_from_question,
    is_table_question,
    matching_figure_indices,
    matching_table_indices,
    rerank_table_first,
    split_to_chunks,
    supplement_answer_with_evidence,
    table_number_from_question,
    validate_answer_against_evidence,
)
from sci_rag_reranking import CrossEncoderReranker, reranker_document_text
from sci_rag_retrieval import (
    BM25Index,
    DocumentRouter,
    RankedItem,
    query_variants,
    reciprocal_rank_fusion,
    tokenize,
)


@dataclass(frozen=True)
class RuntimeConfig:
    """Runtime settings, all overridable through environment variables."""

    embedding_model: str = "BAAI/bge-small-zh-v1.5"
    db_path: str = "./chroma_db"
    deepseek_base_url: str = "https://api.deepseek.com/v1"
    deepseek_model: str = "deepseek-chat"
    retrieval_k: int = 12
    context_k: int = 10
    retrieval_mode: str = "dense"
    document_routing: bool = False
    query_decomposition: bool = False
    parent_window: bool = False
    spatial_figure_evidence: bool = False
    formula_evidence: bool = False
    answer_validation: bool = False
    hybrid_candidate_k: int = 50
    hybrid_rrf_k: int = 60
    reranker_model: str | None = None
    reranker_revision: str | None = None
    reranker_batch_size: int = 8
    reranker_max_length: int = 512
    reranker_device: str = "cpu"
    reranker_rrf_k: int = 60

    @classmethod
    def from_env(cls) -> "RuntimeConfig":
        def positive_int(name: str, default: int) -> int:
            try:
                value = int(os.getenv(name, str(default)))
            except (TypeError, ValueError):
                return default
            return value if value > 0 else default

        def optional_text(name: str, default: str | None = None) -> str | None:
            value = os.getenv(name, default or "").strip()
            return value or None

        retrieval_mode = os.getenv("SCI_RAG_RETRIEVAL_MODE", cls.retrieval_mode).strip().casefold()
        if retrieval_mode not in {"dense", "hybrid"}:
            retrieval_mode = cls.retrieval_mode
        document_routing = os.getenv("SCI_RAG_DOCUMENT_ROUTING", "0").strip().casefold() in {
            "1",
            "true",
            "yes",
            "on",
        }
        query_decomposition = os.getenv(
            "SCI_RAG_QUERY_DECOMPOSITION", "0"
        ).strip().casefold() in {"1", "true", "yes", "on"}
        parent_window = os.getenv("SCI_RAG_PARENT_WINDOW", "0").strip().casefold() in {
            "1",
            "true",
            "yes",
            "on",
        }
        spatial_figure_evidence = os.getenv(
            "SCI_RAG_SPATIAL_FIGURE_EVIDENCE", "0"
        ).strip().casefold() in {"1", "true", "yes", "on"}
        formula_evidence = os.getenv(
            "SCI_RAG_FORMULA_EVIDENCE", "0"
        ).strip().casefold() in {"1", "true", "yes", "on"}
        answer_validation = os.getenv(
            "SCI_RAG_ANSWER_VALIDATION", "0"
        ).strip().casefold() in {"1", "true", "yes", "on"}
        return cls(
            embedding_model=os.getenv("SCI_RAG_EMBEDDING_MODEL", cls.embedding_model),
            db_path=os.getenv("SCI_RAG_DB_PATH", cls.db_path),
            deepseek_base_url=os.getenv("DEEPSEEK_BASE_URL", cls.deepseek_base_url),
            deepseek_model=os.getenv("DEEPSEEK_MODEL", cls.deepseek_model),
            retrieval_k=positive_int("SCI_RAG_RETRIEVAL_K", cls.retrieval_k),
            context_k=positive_int("SCI_RAG_CONTEXT_K", cls.context_k),
            retrieval_mode=retrieval_mode,
            document_routing=document_routing,
            query_decomposition=query_decomposition,
            parent_window=parent_window,
            spatial_figure_evidence=spatial_figure_evidence,
            formula_evidence=formula_evidence,
            answer_validation=answer_validation,
            hybrid_candidate_k=positive_int(
                "SCI_RAG_HYBRID_CANDIDATE_K", cls.hybrid_candidate_k
            ),
            hybrid_rrf_k=positive_int("SCI_RAG_HYBRID_RRF_K", cls.hybrid_rrf_k),
            reranker_model=optional_text("SCI_RAG_RERANKER_MODEL"),
            reranker_revision=optional_text("SCI_RAG_RERANKER_REVISION"),
            reranker_batch_size=positive_int(
                "SCI_RAG_RERANKER_BATCH_SIZE", cls.reranker_batch_size
            ),
            reranker_max_length=positive_int(
                "SCI_RAG_RERANKER_MAX_LENGTH", cls.reranker_max_length
            ),
            reranker_device=os.getenv(
                "SCI_RAG_RERANKER_DEVICE", cls.reranker_device
            ).strip()
            or cls.reranker_device,
            reranker_rrf_k=positive_int(
                "SCI_RAG_RERANKER_RRF_K", cls.reranker_rrf_k
            ),
        )


@dataclass
class LexicalSnapshot:
    """Cached collection text used by optional lexical and source routing."""

    collection_count: int
    ids: list[str]
    texts: list[str]
    metadatas: list[dict[str, Any]]
    index: BM25Index
    router: DocumentRouter | None = None


class Runtime:
    """Explicitly initialized model/API/database resources."""

    def __init__(
        self,
        config: RuntimeConfig,
        client: Any,
        embedding_model: Any,
        collection: Any,
        reranker: Any | None = None,
    ):
        if reranker is not None and config.retrieval_mode != "hybrid":
            raise ValueError("cross-encoder reranker 只能与 hybrid 检索一起启用")
        self.config = config
        self.client = client
        self.embedding_model = embedding_model
        self.collection = collection
        self.reranker = reranker
        self._lexical_snapshot: LexicalSnapshot | None = None

    def invalidate_lexical_index(self) -> None:
        self._lexical_snapshot = None


_runtime: Runtime | None = None


def create_runtime(config: RuntimeConfig | None = None) -> Runtime:
    """Initialize external resources exactly once per caller-owned runtime."""

    from dotenv import load_dotenv

    load_dotenv()
    config = config or RuntimeConfig.from_env()
    if config.reranker_model and config.retrieval_mode != "hybrid":
        raise ValueError("SCI_RAG_RERANKER_MODEL 需要 SCI_RAG_RETRIEVAL_MODE=hybrid")
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        raise ValueError("请在 .env 文件中设置 DEEPSEEK_API_KEY")

    from openai import OpenAI
    from sentence_transformers import SentenceTransformer
    import chromadb

    client = OpenAI(api_key=api_key, base_url=config.deepseek_base_url)
    embedding_model = SentenceTransformer(config.embedding_model)
    reranker = None
    if config.reranker_model:
        os.environ.setdefault("HF_HUB_OFFLINE", "1")
        reranker = CrossEncoderReranker(
            config.reranker_model,
            revision=config.reranker_revision,
            batch_size=config.reranker_batch_size,
            max_length=config.reranker_max_length,
            device=config.reranker_device,
            local_files_only=True,
        )
    chroma_client = chromadb.PersistentClient(path=config.db_path)
    collection = chroma_client.get_or_create_collection(
        name="knowledge_base",
        metadata={"hnsw:space": "cosine"},
    )
    return Runtime(config, client, embedding_model, collection, reranker=reranker)


def get_runtime() -> Runtime:
    """Lazily initialize the default runtime for legacy callers."""

    global _runtime
    if _runtime is None:
        _runtime = create_runtime()
    return _runtime


def _docx_to_markdown(file_path: str) -> str:
    """Read DOCX with the declared ``python-docx`` dependency."""

    from docx import Document as WordDocument

    document = WordDocument(file_path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        rows = [[cell.text.replace("\n", " ").strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        table_lines = ["|" + "|".join(rows[0]) + "|", "|" + "|".join("---" for _ in rows[0]) + "|"]
        table_lines.extend("|" + "|".join(row) + "|" for row in rows[1:])
        parts.append("\n".join(table_lines))
    return "\n\n".join(parts)


def load_and_split_document(
    file_path: str,
    include_spatial_figures: bool = False,
) -> list[Chunk]:
    """Load PDF/TXT/DOCX and return canonical, page-aware chunks.

    Spatial figure evidence is experimental and opt-in. It reads coordinates
    already present in a born-digital PDF text layer; it neither persists nor
    interprets image pixels.
    """

    source = os.path.basename(file_path)
    suffix = Path(file_path).suffix.lower()
    documents: list[Chunk] = []

    if suffix == ".pdf":
        try:
            import pymupdf
            import pymupdf4llm
        except ImportError as exc:
            raise ImportError(
                "缺少 PDF 依赖，请安装 pymupdf4llm、pymupdf 和 pillow。"
            ) from exc

        document = pymupdf.open(file_path)
        try:
            page_chunks = pymupdf4llm.to_markdown(
                document,
                filename=source,
                page_chunks=True,
                table_output="markdown",
                write_images=False,
                embed_images=False,
            )
            if isinstance(page_chunks, list):
                for page_number, page_chunk in enumerate(page_chunks, start=1):
                    text = str(page_chunk.get("text", ""))
                    if text.strip():
                        documents.append(
                            Chunk(
                                page_content=text,
                                metadata={"source": source, "page": page_number},
                            )
                        )
            else:
                documents.append(Chunk(str(page_chunks), {"source": source}))
            if include_spatial_figures:
                for page_number, page in enumerate(document, start=1):
                    documents.extend(
                        extract_spatial_figure_chunks(
                            page.get_text("blocks", sort=True),
                            source,
                            page_number,
                            float(page.rect.width),
                            float(page.rect.height),
                        )
                    )
        finally:
            document.close()
    elif suffix == ".txt":
        documents.append(
            Chunk(Path(file_path).read_text(encoding="utf-8"), {"source": source, "page": 1})
        )
    elif suffix == ".docx":
        documents.append(Chunk(_docx_to_markdown(file_path), {"source": source, "page": 1}))
    else:
        raise ValueError("不支持的文件类型（仅支持 .pdf / .txt / .docx）")

    return split_to_chunks(documents, source)


def _metadata_for_chroma(metadata: dict[str, Any]) -> dict[str, Any]:
    """Remove unsupported ``None`` values before writing Chroma metadata."""

    return {key: value for key, value in metadata.items() if value is not None}


def _sha256(file_path: str) -> str:
    digest = hashlib.sha256()
    with open(file_path, "rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def add_document_to_db(file_path: str, runtime: Runtime | None = None) -> str:
    runtime = runtime or get_runtime()
    chunks = load_and_split_document(
        file_path,
        include_spatial_figures=runtime.config.spatial_figure_evidence,
    )
    document_hash = _sha256(file_path)
    for index, chunk in enumerate(chunks):
        text = chunk.page_content
        metadata = dict(chunk.metadata)
        metadata.update(
            {
                "source": os.path.basename(file_path),
                "document_sha256": document_hash,
                "chunk_index": index,
            }
        )
        stable_id = hashlib.sha256(
            f"{document_hash}:{index}:{metadata.get('type', 'text')}:{text}".encode("utf-8")
        ).hexdigest()
        metadata["chunk_id"] = stable_id
        embedding = runtime.embedding_model.encode(text).tolist()
        runtime.collection.upsert(
            ids=[stable_id],
            embeddings=[embedding],
            documents=[text],
            metadatas=[_metadata_for_chroma(metadata)],
        )
    runtime.invalidate_lexical_index()
    return f"✅ 成功添加 {len(chunks)} 个文本块到知识库，当前共 {runtime.collection.count()} 个。"


def upload_file(file: Any, runtime: Runtime | None = None) -> str:
    if file is None:
        return "请选择一个文件"
    return add_document_to_db(file.name, runtime=runtime)


def _merge_results(
    dense: dict[str, Any],
    tables: dict[str, Any] | None,
    additional_results: list[dict[str, Any]] | None = None,
) -> tuple[list[str], list[str], list[dict[str, Any]]]:
    texts: list[str] = []
    ids: list[str] = []
    metas: list[dict[str, Any]] = []
    seen: set[str] = set()

    def add(result: dict[str, Any]) -> None:
        result_ids = (result.get("ids") or [[]])[0]
        result_docs = (result.get("documents") or [[]])[0]
        result_metas = (result.get("metadatas") or [[]])[0]
        for index, doc_id in enumerate(result_ids):
            if doc_id in seen:
                continue
            seen.add(doc_id)
            ids.append(doc_id)
            texts.append(result_docs[index] if index < len(result_docs) else "")
            metas.append(result_metas[index] if index < len(result_metas) else {})

    for result in [*(additional_results or []), dense]:
        add(result)
    if tables:
        # Collection.get() returns flat lists, unlike query()'s nested lists.
        flat_ids = tables.get("ids") or []
        flat_docs = tables.get("documents") or []
        flat_metas = tables.get("metadatas") or []
        add({
            "ids": [flat_ids],
            "documents": [flat_docs],
            "metadatas": [flat_metas],
        })
    return texts, ids, metas


def _flat_result_values(result: dict[str, Any], key: str) -> list[Any]:
    """Read Chroma get/query values while tolerating simple test doubles."""

    values = result.get(key) or []
    if values and isinstance(values[0], list):
        values = values[0]
    return list(values)


def _lexical_search_text(text: str, metadata: dict[str, Any]) -> str:
    return "\n".join(
        [
            text,
            str(metadata.get("table_caption", "")),
            str(metadata.get("headers", "")),
            str(metadata.get("source", "")),
        ]
    )


_TABLE_UNIT_HINT_RE = re.compile(
    r"(?:[×x]\s*10|%|Å|kcal(?:\s*/\s*mol)?|s\s*/\s*iter|\bunits?\b)",
    re.IGNORECASE,
)


def _table_caption_unit_note(metadata: dict[str, Any]) -> str:
    """Return a concise table-note suffix when the caption declares units.

    Deterministic table-cell lookup otherwise returns only the row/column value.
    Captions are the authoritative place where some scientific tables declare a
    shared scale (for example ``×10^-2``), so preserve that note without
    inventing or converting a value.  Captions without an explicit unit hint
    are intentionally omitted to keep existing concise answers unchanged.
    """

    caption = str(metadata.get("table_caption", "") or "")
    if not caption:
        return ""
    plain = html.unescape(re.sub(r"<[^>]*>", " ", caption))
    plain = re.sub(r"\s+", " ", plain).strip()
    if not plain or not _TABLE_UNIT_HINT_RE.search(plain):
        return ""
    body = re.sub(
        r"^(?:extended\s+data\s+)?table\s*\d+\s*[:.]?\s*",
        "",
        plain,
        flags=re.IGNORECASE,
    ).strip()
    return f"（表注：{body or plain}）"


def _fuse_dense_results(
    results: list[dict[str, Any]],
    limit: int,
    rrf_k: int,
) -> dict[str, Any]:
    """Fuse dense rankings produced for deterministic query variants."""

    if not results:
        return {"ids": [[]], "documents": [[]], "metadatas": [[]]}
    id_lists = [
        [str(value) for value in _flat_result_values(result, "ids")]
        for result in results
    ]
    fused = reciprocal_rank_fusion(id_lists, rrf_k=rrf_k, limit=limit)
    payloads: dict[str, tuple[str, dict[str, Any]]] = {}
    for result in results:
        ids = [str(value) for value in _flat_result_values(result, "ids")]
        texts = _flat_result_values(result, "documents")
        metas = _flat_result_values(result, "metadatas")
        for index, doc_id in enumerate(ids):
            if doc_id in payloads:
                continue
            metadata = metas[index] if index < len(metas) and isinstance(metas[index], dict) else {}
            text = str(texts[index]) if index < len(texts) else ""
            payloads[doc_id] = (text, dict(metadata))
    ids: list[str] = []
    texts: list[str] = []
    metadatas: list[dict[str, Any]] = []
    for item in fused:
        doc_id = str(item.key)
        payload = payloads.get(doc_id)
        if payload is None:
            continue
        ids.append(doc_id)
        texts.append(payload[0])
        metadatas.append(payload[1])
    return {"ids": [ids], "documents": [texts], "metadatas": [metadatas]}


def _get_lexical_snapshot(runtime: Runtime) -> LexicalSnapshot:
    """Build or reuse lexical indexes for the current Chroma collection."""

    collection_count = runtime.collection.count()
    cached = runtime._lexical_snapshot
    if cached is not None and cached.collection_count == collection_count:
        return cached

    result = runtime.collection.get(include=["documents", "metadatas"])
    raw_ids = _flat_result_values(result, "ids")
    raw_texts = _flat_result_values(result, "documents")
    raw_metas = _flat_result_values(result, "metadatas")
    ids = [str(value) for value in raw_ids]
    texts = [str(raw_texts[index]) if index < len(raw_texts) else "" for index in range(len(ids))]
    metadatas = [
        dict(raw_metas[index]) if index < len(raw_metas) and isinstance(raw_metas[index], dict) else {}
        for index in range(len(ids))
    ]
    search_documents = [
        _lexical_search_text(text, metadata) for text, metadata in zip(texts, metadatas)
    ]
    source_profiles: dict[str, list[str]] = {}
    for text, metadata in zip(texts, metadatas):
        if (
            runtime.config.spatial_figure_evidence
            and metadata.get("type") == "figure"
        ):
            continue
        source = str(metadata.get("source", "")).strip()
        if source:
            source_profiles.setdefault(source, []).append(
                _lexical_search_text(text, metadata)
            )
    router = None
    if len(source_profiles) >= 1:
        router = DocumentRouter(
            source_profiles.keys(),
            ("\n".join(parts) for parts in source_profiles.values()),
        )
    snapshot = LexicalSnapshot(
        collection_count=collection_count,
        ids=ids,
        texts=texts,
        metadatas=metadatas,
        index=BM25Index(search_documents),
        router=router,
    )
    runtime._lexical_snapshot = snapshot
    return snapshot


def _hybrid_fused_result(
    question: str,
    dense: dict[str, Any],
    runtime: Runtime,
    candidate_k: int,
    source_filter: str | None = None,
    lexical_queries: list[str] | None = None,
) -> dict[str, Any]:
    """Fuse Chroma dense results with a cached lexical ranking using RRF."""

    dense_ids = [str(value) for value in _flat_result_values(dense, "ids")]
    dense_texts = _flat_result_values(dense, "documents")
    dense_metas = _flat_result_values(dense, "metadatas")
    dense_by_id = {
        doc_id: (
            str(dense_texts[index]) if index < len(dense_texts) else "",
            dict(dense_metas[index])
            if index < len(dense_metas) and isinstance(dense_metas[index], dict)
            else {},
        )
        for index, doc_id in enumerate(dense_ids)
    }

    snapshot = _get_lexical_snapshot(runtime)
    lexical_indices = None
    if source_filter or runtime.config.spatial_figure_evidence:
        lexical_indices = [
            index
            for index, metadata in enumerate(snapshot.metadatas)
            if (
                (not source_filter or str(metadata.get("source", "")) == source_filter)
                and (
                    not runtime.config.spatial_figure_evidence
                    or metadata.get("type") != "figure"
                )
            )
        ]
    lexical_rankings: list[list[str]] = []
    for lexical_query in lexical_queries or [question]:
        if not snapshot.index.has_lexical_signal(lexical_query):
            continue
        lexical_rankings.append(
            [
                snapshot.ids[int(item.key)]
                for item in snapshot.index.retrieve(
                    lexical_query, candidate_k, indices=lexical_indices
                )
            ]
        )
    lexical_ids = (
        [
            str(item.key)
            for item in reciprocal_rank_fusion(
                lexical_rankings,
                rrf_k=runtime.config.hybrid_rrf_k,
                limit=candidate_k,
            )
        ]
        if lexical_rankings
        else []
    )
    snapshot_by_id = {
        doc_id: (snapshot.texts[index], snapshot.metadatas[index])
        for index, doc_id in enumerate(snapshot.ids)
    }
    fused = reciprocal_rank_fusion(
        [dense_ids, lexical_ids],
        rrf_k=runtime.config.hybrid_rrf_k,
        limit=candidate_k,
    )

    ids: list[str] = []
    texts: list[str] = []
    metadatas: list[dict[str, Any]] = []
    for item in fused:
        doc_id = str(item.key)
        payload = dense_by_id.get(doc_id) or snapshot_by_id.get(doc_id)
        if payload is None:
            continue
        text, metadata = payload
        ids.append(doc_id)
        texts.append(text)
        metadatas.append(metadata)
    return {"ids": [ids], "documents": [texts], "metadatas": [metadatas]}


def _cross_encoder_reranked_result(
    question: str,
    result: dict[str, Any],
    runtime: Runtime,
) -> dict[str, Any]:
    """Rerank Hybrid candidates, then conservatively fuse the original order."""

    if runtime.reranker is None:
        return result
    ids = [str(value) for value in _flat_result_values(result, "ids")]
    raw_texts = _flat_result_values(result, "documents")
    raw_metas = _flat_result_values(result, "metadatas")
    texts = [str(raw_texts[index]) if index < len(raw_texts) else "" for index in range(len(ids))]
    metadatas = [
        dict(raw_metas[index])
        if index < len(raw_metas) and isinstance(raw_metas[index], dict)
        else {}
        for index in range(len(ids))
    ]
    candidates = [RankedItem(index, 0.0) for index in range(len(ids))]
    passages = [
        reranker_document_text(text, metadata)
        for text, metadata in zip(texts, metadatas)
    ]
    reranked = runtime.reranker.rerank(question, candidates, passages).ranked
    fused = reciprocal_rank_fusion(
        [reranked, candidates],
        rrf_k=runtime.config.reranker_rrf_k,
        limit=len(candidates),
    )
    order = [int(item.key) for item in fused]
    return {
        "ids": [[ids[index] for index in order]],
        "documents": [[texts[index] for index in order]],
        "metadatas": [[metadatas[index] for index in order]],
    }


_COMPOSITE_FACT_CUE_RE = re.compile(
    r"多少|哪些|如何|管道|步骤|阶段|以及|并且|同时|与|和|"
    r"(?:[一二三四五六七八九十0-9]+种|多个|两种|若干).{0,20}(?:什么|哪些|分别)|"
    r"\b(?:what|which|how|and|pipeline|dataset)\b",
    re.IGNORECASE,
)
_LISTED_FACT_QUESTION_RE = re.compile(
    r"(?:[一二三四五六七八九十0-9]+种|多个|两种|若干).{0,20}(?:什么|哪些|分别)",
    re.IGNORECASE,
)
_REFERENCE_HEADER_RE = re.compile(
    r"(?:references?|bibliography|参考文献)", re.IGNORECASE
)
_PICTURE_TEXT_MARKER_RE = re.compile(
    r"<!--\s*(?:start|end) of picture text\s*-->|<img\b",
    re.IGNORECASE,
)
_SECTION_QUERY_ALIASES = {
    "数据集": ("dataset", "data"),
    "显式推理": ("explicit-reasoning", "reasoning"),
    "推理": ("reasoning",),
    "强化学习": ("reinforcement", "rl"),
    "训练": ("training", "train"),
    "实验配置": ("experimental", "setup", "configuration", "configurations"),
    "配置": ("configuration", "configurations", "setup"),
    "奖励": ("reward",),
    "管道": ("pipeline",),
}
MAX_SECTION_EXPANSION_CHUNKS = 6
MAX_PARENT_WINDOW_ANCHORS = 2


def _is_picture_text_chunk(text: Any) -> bool:
    """Identify pymupdf4llm's OCR/image-text blocks for window isolation."""

    return bool(_PICTURE_TEXT_MARKER_RE.search(str(text or "")))


def _is_composite_fact_question(question: str) -> bool:
    """Detect questions likely to require evidence from multiple chunks."""

    text = str(question or "")
    matches = _COMPOSITE_FACT_CUE_RE.findall(text)
    return len(matches) >= 2 or bool(_LISTED_FACT_QUESTION_RE.search(text))


def _section_query_terms(question: str) -> set[str]:
    terms = {token for token in tokenize(question) if len(token) >= 3}
    normalized = str(question or "").casefold()
    for phrase, aliases in _SECTION_QUERY_ALIASES.items():
        if phrase in normalized:
            terms.update(aliases)
    return terms


def _header_match_score(header: str, query_terms: set[str]) -> int:
    """Score the deepest heading more heavily than inherited parent headings."""

    parts = [part.strip() for part in str(header).split(">") if part.strip()]
    if not parts:
        return 0
    deepest = set(tokenize(parts[-1]))
    score = len(deepest & query_terms) * 3
    if len(parts) > 1:
        score += len(set(tokenize(parts[-2])) & query_terms)
    return score


def _section_expansion_result(
    question: str,
    base_result: dict[str, Any],
    runtime: Runtime,
    route: Any | None = None,
) -> dict[str, Any] | None:
    """Add same-section chunks for composite questions without cross-paper mixing.

    PDF-to-Markdown chunkers keep the heading path in ``metadata['headers']``.
    A multi-fact question can retrieve a section's overview while missing the
    immediately following chunk that contains a threshold or tool name. This
    bounded expansion reads the existing collection, stays within the source
    of the highest-ranked candidate, selects the strongest matching header, and
    adds at most a small neighborhood of chunks before the final context cap.
    It never invents text or facts or mixes papers merely because their headings
    use the same generic terms.
    """

    if not _is_composite_fact_question(question):
        return None
    base_ids = [str(value) for value in _flat_result_values(base_result, "ids")]
    base_metas = _flat_result_values(base_result, "metadatas")
    anchor_ids = base_ids[: runtime.config.context_k]
    anchor_metas = base_metas[: runtime.config.context_k]
    anchor_sources = {
        str(meta.get("source", ""))
        for meta in anchor_metas
        if isinstance(meta, dict) and meta.get("source")
    }
    if len(anchor_sources) > 1:
        return None
    base_source = ""
    for meta in anchor_metas:
        if isinstance(meta, dict) and meta.get("source"):
            base_source = str(meta["source"])
            break
    if not base_source:
        return None
    all_result = runtime.collection.get(include=["documents", "metadatas"])
    all_texts = _flat_result_values(all_result, "documents")
    all_metas = _flat_result_values(all_result, "metadatas")
    corpus_sources = {
        str(meta.get("source", ""))
        for meta in all_metas
        if isinstance(meta, dict) and meta.get("source")
    }
    route_source = str(getattr(route, "document_id", "") or "").strip()
    if len(corpus_sources) > 1 and not route_source:
        return None
    selected_source = route_source or base_source
    if selected_source not in corpus_sources:
        return None
    query_terms = _section_query_terms(question)
    if not query_terms:
        return None

    header_rows: list[tuple[int, int, str, str]] = []
    for index, raw_meta in enumerate(all_metas):
        metadata = raw_meta if isinstance(raw_meta, dict) else {}
        source = str(metadata.get("source", ""))
        header = str(metadata.get("headers", ""))
        if source != selected_source or not header:
            continue
        score = _header_match_score(header, query_terms)
        if score:
            header_rows.append((score, index, source, header))
    if not header_rows:
        return None

    best_score = max(row[0] for row in header_rows)
    selected_headers = {(row[2], row[3]) for row in header_rows if row[0] == best_score}
    all_ids = _flat_result_values(all_result, "ids")
    all_id_to_index = {str(doc_id): index for index, doc_id in enumerate(all_ids)}
    anchors = [
        all_id_to_index[doc_id]
        for doc_id in anchor_ids
        if doc_id in all_id_to_index
        and (
            str(
                all_metas[all_id_to_index[doc_id]].get("source", "")
                if isinstance(all_metas[all_id_to_index[doc_id]], dict)
                else ""
            ),
            str(
                all_metas[all_id_to_index[doc_id]].get("headers", "")
                if isinstance(all_metas[all_id_to_index[doc_id]], dict)
                else ""
            ),
        ) in selected_headers
    ]
    if not anchors:
        return None
    selected_indices = [
        index
        for index, raw_meta in enumerate(all_metas)
        if (
            str(raw_meta.get("source", "")) if isinstance(raw_meta, dict) else "",
            str(raw_meta.get("headers", "")) if isinstance(raw_meta, dict) else "",
        ) in selected_headers
        and index < len(all_texts)
        and index < len(all_ids)
    ]
    selected_indices.sort(
        key=lambda index: (
            min((abs(index - anchor) for anchor in anchors), default=index),
            index,
        )
    )
    selected_ids: list[str] = []
    selected_docs: list[str] = []
    selected_metas: list[dict[str, Any]] = []
    for index in selected_indices[:MAX_SECTION_EXPANSION_CHUNKS]:
        metadata = all_metas[index] if isinstance(all_metas[index], dict) else {}
        doc_id = str(all_ids[index])
        selected_ids.append(doc_id)
        selected_docs.append(str(all_texts[index]))
        selected_metas.append(dict(metadata))

    if not selected_ids:
        return None
    return {
        "ids": [selected_ids],
        "documents": [selected_docs],
        "metadatas": [selected_metas],
    }


def _parent_window_contexts(
    texts: list[str],
    ids: list[str],
    metas: list[dict[str, Any]],
    runtime: Runtime,
) -> tuple[list[str], list[dict[str, Any]]]:
    """Attach same-page neighbors to top text anchors without adding slots.

    Uploaded chunks carry a source-local ``chunk_index``.  The optional parent
    window uses that stable sequence rather than Collection.get() ordering,
    skips tables/references and already-selected contexts, and records every
    contributing chunk ID in returned metadata.  The returned text is therefore
    exactly what generation receives while citations retain the anchor page.
    """

    if not runtime.config.parent_window or not texts:
        return list(texts), [dict(meta) for meta in metas]
    snapshot = _get_lexical_snapshot(runtime)
    sequence: dict[tuple[str, int], tuple[str, str, dict[str, Any]]] = {}
    for doc_id, text, raw_meta in zip(
        snapshot.ids, snapshot.texts, snapshot.metadatas
    ):
        metadata = raw_meta if isinstance(raw_meta, dict) else {}
        source = str(metadata.get("source", ""))
        try:
            chunk_index = int(metadata.get("chunk_index"))
        except (TypeError, ValueError):
            continue
        if source:
            sequence[(source, chunk_index)] = (doc_id, text, metadata)

    selected_ids = {str(doc_id) for doc_id in ids}
    effective_texts = list(texts)
    effective_metas = [dict(meta) for meta in metas]
    for position in range(min(MAX_PARENT_WINDOW_ANCHORS, len(effective_texts))):
        metadata = effective_metas[position]
        if metadata.get("type", "text") != "text":
            continue
        # Figure/OCR blocks often contain several unrelated axes and sample
        # counts.  Expanding them with adjacent prose increases numeric
        # ambiguity, so retain the ranked block but do not create a window.
        if _is_picture_text_chunk(effective_texts[position]):
            continue
        source = str(metadata.get("source", ""))
        page = metadata.get("page")
        header = str(metadata.get("headers", ""))
        try:
            anchor_chunk_index = int(metadata.get("chunk_index"))
        except (TypeError, ValueError):
            continue
        if not source or page is None or _REFERENCE_HEADER_RE.search(header):
            continue
        anchor_record = sequence.get((source, anchor_chunk_index))
        if anchor_record is None:
            continue
        included = [(anchor_chunk_index, *anchor_record)]
        for neighbor_index in (anchor_chunk_index - 1, anchor_chunk_index + 1):
            record = sequence.get((source, neighbor_index))
            if record is None:
                continue
            neighbor_id, neighbor_text, neighbor_meta = record
            if str(neighbor_id) in selected_ids:
                continue
            if neighbor_meta.get("type", "text") != "text":
                continue
            if _is_picture_text_chunk(neighbor_text):
                continue
            if neighbor_meta.get("page") != page:
                continue
            if _REFERENCE_HEADER_RE.search(str(neighbor_meta.get("headers", ""))):
                continue
            included.append(
                (neighbor_index, str(neighbor_id), str(neighbor_text), neighbor_meta)
            )
        included.sort(key=lambda row: row[0])
        if len(included) <= 1:
            continue
        effective_texts[position] = "\n\n".join(row[2] for row in included)
        effective_metas[position]["window_chunk_indices"] = [
            row[0] for row in included
        ]
        effective_metas[position]["window_chunk_ids"] = [row[1] for row in included]
        effective_metas[position]["window_added_character_count"] = sum(
            len(row[2]) for row in included if row[1] != str(ids[position])
        )
    return effective_texts, effective_metas


def query_knowledge(
    message: str,
    history: Any = None,
    return_contexts: bool = True,
    runtime: Runtime | None = None,
) -> str | dict[str, Any]:
    """Retrieve evidence and generate an answer.

    When ``return_contexts`` is true, ``contexts`` is exactly the list joined
    into the generation prompt. IDs and metadata are returned for evaluation.
    ``history`` is accepted for Gradio compatibility but unused in this
    single-turn baseline.
    """

    runtime = runtime or get_runtime()
    if not message or not message.strip():
        result = {"answer": "请输入一个问题。", "contexts": [], "context_ids": [], "context_metadatas": []}
        return result if return_contexts else result["answer"]
    if runtime.collection.count() == 0:
        answer = "📚 知识库为空，请先上传文档。"
        return {"answer": answer, "contexts": [], "context_ids": [], "context_metadatas": []} if return_contexts else answer

    variants = (
        query_variants(message)
        if runtime.config.query_decomposition
        else [message]
    )
    if not variants:
        variants = [message]
    configured_candidate_k = (
        runtime.config.hybrid_candidate_k
        if runtime.config.retrieval_mode == "hybrid"
        else runtime.config.retrieval_k
    )
    candidate_k = min(configured_candidate_k, runtime.collection.count())
    route = None
    if runtime.config.document_routing:
        route_snapshot = _get_lexical_snapshot(runtime)
        route = route_snapshot.router.route(message) if route_snapshot.router else None
    dense_results: list[dict[str, Any]] = []
    for variant in variants:
        question_embedding = runtime.embedding_model.encode(variant).tolist()
        query_kwargs = {
            "query_embeddings": [question_embedding],
            "n_results": candidate_k,
            "include": ["documents", "metadatas", "distances"],
        }
        if runtime.config.spatial_figure_evidence:
            # Figure blocks are an explicit Figure N evidence channel. Keeping
            # them out of ordinary dense/Hybrid candidates prevents extra
            # diagram labels from changing unrelated questions' rankings.
            query_conditions: list[dict[str, Any]] = [
                {"type": {"$ne": "figure"}}
            ]
            if route is not None:
                query_conditions.append({"source": {"$eq": route.document_id}})
            query_kwargs["where"] = (
                query_conditions[0]
                if len(query_conditions) == 1
                else {"$and": query_conditions}
            )
        elif route is not None:
            # ``source`` is written for every uploaded chunk.  The filter is
            # only applied after the conservative lexical router found one
            # unique source; ambiguous questions intentionally keep the global
            # search for every variant.
            query_kwargs["where"] = {"source": route.document_id}
        dense_results.append(runtime.collection.query(**query_kwargs))
    dense = _fuse_dense_results(
        dense_results,
        candidate_k,
        runtime.config.hybrid_rrf_k,
    )
    if runtime.config.retrieval_mode == "hybrid":
        dense = _hybrid_fused_result(
            message,
            dense,
            runtime,
            candidate_k,
            source_filter=route.document_id if route is not None else None,
            lexical_queries=variants,
        )
        dense = _cross_encoder_reranked_result(message, dense, runtime)
    section_result = _section_expansion_result(message, dense, runtime, route=route)
    formula_result = None
    if runtime.config.formula_evidence:
        formula_snapshot = _get_lexical_snapshot(runtime)
        source_ids = {
            str(metadata.get("source", "")).strip()
            for metadata in formula_snapshot.metadatas
            if isinstance(metadata, dict) and str(metadata.get("source", "")).strip()
        }
        allowed_formula_indices: list[int] | None
        if route is not None:
            allowed_formula_indices = [
                index
                for index in range(len(formula_snapshot.texts))
                if str(formula_snapshot.metadatas[index].get("source", ""))
                == route.document_id
            ]
        elif len(source_ids) == 1:
            # A single-source collection is safe for this opt-in lexical aid.
            allowed_formula_indices = list(range(len(formula_snapshot.texts)))
        else:
            # Do not let an unqualified multi-paper question borrow an
            # equation from an unrelated source. The normal retriever remains
            # responsible for resolving the ambiguous query.
            allowed_formula_indices = []
        formula_indices = formula_evidence_indices(
            message,
            formula_snapshot.texts,
            formula_snapshot.metadatas,
            max_results=8,
            allowed_indices=allowed_formula_indices,
        )
        if formula_indices:
            formula_result = {
                "ids": [[formula_snapshot.ids[index] for index in formula_indices]],
                "documents": [[formula_snapshot.texts[index] for index in formula_indices]],
                "metadatas": [[
                    {**formula_snapshot.metadatas[index], "formula_evidence": True}
                    for index in formula_indices
                ]],
            }
    figure_result = None
    explicit_figure_reference = figure_reference_from_question(message)
    if runtime.config.spatial_figure_evidence and explicit_figure_reference is not None:
        figure_kind, figure_number = explicit_figure_reference
        figure_conditions: list[dict[str, Any]] = [
            {"type": {"$eq": "figure"}},
            {"figure_kind": {"$eq": figure_kind}},
            {"figure_number": {"$eq": figure_number}},
        ]
        if route is not None:
            figure_conditions.append({"source": {"$eq": route.document_id}})
        raw_figure_result = runtime.collection.get(
            where={"$and": figure_conditions},
            include=["documents", "metadatas"],
        )
        # ``collection.get`` is flat whereas normal query results are nested.
        # Wrap it so the common merge path can place exact Figure N evidence
        # before dense candidates without special-casing IDs or metadata.
        figure_result = {
            "ids": [raw_figure_result.get("ids") or []],
            "documents": [raw_figure_result.get("documents") or []],
            "metadatas": [raw_figure_result.get("metadatas") or []],
        }
    table_results = None
    if is_table_question(message):
        table_where: dict[str, Any] = {"type": "table"}
        if route is not None:
            # A source route already narrowed the dense/lexical candidate pool.
            # Keep the deterministic table scan in that same scope; otherwise
            # Table N from a different uploaded paper could satisfy the same
            # row/column names before the routed paper is inspected.
            table_where = {
                "$and": [
                    {"type": {"$eq": "table"}},
                    {"source": {"$eq": route.document_id}},
                ]
            }
        table_results = runtime.collection.get(
            where=table_where,
            include=["documents", "metadatas"],
        )
    retrieved_texts, retrieved_ids, retrieved_metas = _merge_results(
        dense,
        table_results,
        [
            result
            for result in (figure_result, formula_result, section_result)
            if result is not None
        ]
        or None,
    )
    if not retrieved_texts:
        answer = "未找到相关内容，请换个问法。"
        return {"answer": answer, "contexts": [], "context_ids": [], "context_metadatas": []} if return_contexts else answer

    explicit_table_number = table_number_from_question(message)
    matching_tables = matching_table_indices(message, retrieved_texts, retrieved_metas)
    if explicit_table_number is not None and not matching_tables:
        answer = f"在 Table {explicit_table_number} 中未找到可用的结构化表格，不能用其他表格替代。"
        result = {
            "answer": answer,
            "contexts": [],
            "context_ids": [],
            "context_metadatas": [],
        }
        return result if return_contexts else answer

    order, note, filtered_texts = rerank_table_first(message, retrieved_texts, retrieved_metas)
    if explicit_figure_reference is not None and runtime.config.spatial_figure_evidence:
        matching_figures = matching_figure_indices(
            message,
            filtered_texts,
            retrieved_metas,
        )
        if matching_figures:
            matching_set = set(matching_figures)
            order = matching_figures + [
                index
                for index in order
                if index not in matching_set
                and retrieved_metas[index].get("type") != "figure"
                and not _is_picture_text_chunk(filtered_texts[index])
            ]
        else:
            figure_kind, figure_number = explicit_figure_reference
            label = (
                f"Extended Data Figure {figure_number}"
                if figure_kind == "extended_data_figure"
                else f"Figure {figure_number}"
            )
            figure_note = f"未找到 {label} 的坐标化文字证据，已回退普通文本检索。"
            note = f"{note} {figure_note}".strip()

    # A question that names a table, row, and column is a deterministic cell
    # lookup.  Answer it from the parsed table rather than asking a language
    # model to choose between similarly worded Table 1/Table 2 narratives.
    cell_match = find_table_cell_in_chunks(message, filtered_texts, retrieved_metas)
    if cell_match is not None:
        cell_index, cell = cell_match
        table_number = cell["table_number"] or explicit_table_number or "?"
        if "rows" in cell:
            row_texts = []
            for row in cell["rows"]:
                values = "；".join(
                    f"{item['column']}={item['value']}"
                    for item in row.get("values", [])
                )
                row_texts.append(f"{row['row']}：{values}")
            value_text = "；".join(row_texts)
            cell_context = (
                f"Table {table_number} 结构化多行：{value_text}\n\n"
                f"{filtered_texts[cell_index]}"
            )
            answer = f"根据 Table {table_number} 中相关行，列值为：{value_text}。"
        elif "values" in cell:
            value_text = "；".join(
                f"{item['column']}={item['value']}"
                for item in cell["values"]
            )
            cell_context = (
                f"Table {table_number} 结构化行：行={cell['row']}；{value_text}\n\n"
                f"{filtered_texts[cell_index]}"
            )
            answer = f"根据 Table {table_number} 中“{cell['row']}”行，相关列值为：{value_text}。"
        else:
            cell_context = (
                f"Table {table_number} 结构化单元格：行={cell['row']}；"
                f"列={cell['column']}；值={cell['value']}。\n\n{filtered_texts[cell_index]}"
            )
            answer = (
                f"根据 Table {table_number} 中“{cell['row']}”行的“{cell['column']}”列，"
                f"数值为 **{cell['value']}**。"
            )
        table_metadata = (
            retrieved_metas[cell_index]
            if cell_index < len(retrieved_metas)
            and isinstance(retrieved_metas[cell_index], dict)
            else {}
        )
        table_caption = str(table_metadata.get("table_caption", "") or "")
        if table_caption:
            # Keep the caption beside the structured evidence so a returned
            # trace retains shared units/scale notes that are not part of a
            # parsed header (for example a table-wide ×10^-2 footnote).
            cell_context = f"{table_caption}\n\n{cell_context}"
        unit_note = _table_caption_unit_note(table_metadata)
        if unit_note and not _TABLE_UNIT_HINT_RE.search(answer):
            answer += unit_note
        ordered_texts = [cell_context]
        ordered_ids = [retrieved_ids[cell_index]]
        ordered_metas = [retrieved_metas[cell_index]]
        if return_contexts:
            return {
                "answer": answer,
                "contexts": ordered_texts,
                "context_ids": ordered_ids,
                "context_metadatas": ordered_metas,
            }
        source = ordered_metas[0].get("source", "未知")
        page = ordered_metas[0].get("page")
        suffix = f"，第 {page} 页" if page else ""
        suffix += f"（Table {table_number}）"
        return answer + f"\n\n📌 **参考来源：**\n- {source}{suffix}"

    order = order[: runtime.config.context_k]
    ordered_texts = [filtered_texts[index] for index in order]
    ordered_ids = [retrieved_ids[index] for index in order]
    ordered_metas = [retrieved_metas[index] for index in order]
    ordered_texts, ordered_metas = _parent_window_contexts(
        ordered_texts,
        ordered_ids,
        ordered_metas,
        runtime,
    )

    context_parts = []
    for index, (text, metadata) in enumerate(zip(ordered_texts, ordered_metas), start=1):
        if metadata.get("type") == "table":
            table_number = metadata.get("table_number")
            table_label = (
                f"[表格，Table {table_number}]"
                if table_number is not None and str(table_number).strip()
                else "[表格]"
            )
            label = f"【片段 {index}】{table_label}"
        elif metadata.get("type") == "figure":
            label = f"【片段 {index}】[图形坐标文字]"
        elif metadata.get("formula_evidence"):
            label = f"【片段 {index}】[公式候选]"
        else:
            label = f"【片段 {index}】"
        context_parts.append(f"{label}\n{text}")
    context = "\n\n---\n\n".join(context_parts)
    evidence_ledger = build_evidence_ledger(
        message,
        ordered_texts,
        ordered_metas,
    )
    ledger_text = ""
    if evidence_ledger:
        ledger_lines = [f"- {line}" for line in evidence_ledger]
        ledger_text = (
            "【事实核对清单】以下内容仅逐字摘自后面的参考片段，不是新增事实；"
            "回答复合问题时请逐项核对其中与问题相关的数字、阈值、工具名和步骤。\n"
            + "\n".join(ledger_lines)
            + "\n\n"
        )
    if note:
        context = f"【检索提示】{note}\n\n{context}"
    user_prompt = f"{ledger_text}【参考资料】\n{context}\n\n【问题】\n{message}"

    try:
        response = runtime.client.chat.completions.create(
            model=runtime.config.deepseek_model,
            messages=[
                {"role": "system", "content": SCIENTIFIC_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.3,
            max_tokens=2048,
        )
        answer = response.choices[0].message.content
    except Exception as exc:
        answer = f"❌ 调用出错：{exc}"

    has_spatial_figure_context = any(
        metadata.get("type") == "figure" for metadata in ordered_metas
    )
    if (
        _is_composite_fact_question(message)
        and not is_table_question(message)
        and not has_spatial_figure_context
    ):
        answer = supplement_answer_with_evidence(answer, message, evidence_ledger)

    answer_validation = None
    if runtime.config.answer_validation:
        answer_validation = validate_answer_against_evidence(
            message,
            answer,
            evidence_ledger,
        )
        if answer_validation.get("status") == "review":
            reasons = "、".join(answer_validation.get("reasons") or [])
            answer += (
                "\n\n⚠️ **证据核对提示**：生成答案可能遗漏参考片段中的高信号内容，"
                f"请人工核对（{reasons or '未分类原因'}）；系统未自动改写或重试。"
            )

    if return_contexts:
        result = {
            "answer": answer,
            "contexts": ordered_texts,
            "context_ids": ordered_ids,
            "context_metadatas": ordered_metas,
        }
        if answer_validation is not None:
            result["answer_validation"] = answer_validation
        return result
    sources = []
    for metadata in ordered_metas:
        source = metadata.get("source", "未知")
        page = metadata.get("page")
        suffix = f"，第 {page} 页" if page else ""
        if metadata.get("type") == "table":
            suffix += f"（{metadata.get('table_id', '表格')}）"
        elif metadata.get("type") == "figure":
            suffix += f"（{metadata.get('figure_label', '图形坐标文字')}）"
        sources.append(f"- {source}{suffix}")
    unique_sources = list(dict.fromkeys(sources))
    return answer + "\n\n📌 **参考来源：**\n" + "\n".join(unique_sources)


SCIENTIFIC_SYSTEM_PROMPT = """你是一个面向科学论文的严谨学术问答引擎（Sci-RAG），职责是从给定的参考片段中抽取事实、数值与实验方法论。必须遵守以下规则：

【强制规则 1：数值必须原样引用并指明出处】
- 若参考文本中存在具体数值，回答时必须原样引用，不得四舍五入、改写或推算。
- 引用数值后必须指明出处，格式为：根据参考片段 [X] 所示。

【强制规则 1A：图形坐标文字不得跨视觉组拼接】
- 标记为“图形坐标文字”的片段只来自 PDF 文字层坐标，不等同于图片识别。
- 只有标签和值的水平 x 范围重叠时，才可把它们视为同一视觉组；不得把相邻子图、柱或类别中的数值移接到另一标签。
- 如果图中信息只存在于像素而未出现在文字层，必须说明参考片段不足，不能猜测。

【强制规则 2：趋势判断必须有明确对比依据】
- 若问题涉及趋势判断，必须确认参考文本有明确对比依据；没有依据时必须回复“资料未提供该趋势的明确依据，无法推测。”

【强制规则 3：实验步骤按时间顺序重组】
- 若回答涉及实验步骤或方法流程，请按“第一、第二、第三”的逻辑重组叙述，不得调换核心操作顺序或省略中间步骤。

【强制规则 4：实体联合约束】
- 若问题同时指定表格编号和行/列实体名称，必须将两者视为联合约束条件。
- 若在指定表格中找不到实体名称，回复：“在 Table [编号] 中未找到 [实体名称] 的条目。”
- 严禁跨表取数。

【强制规则 5：结构化单元格证据】
- 参考资料中若出现“结构化表格单元格”行，该行是从指定表格的真实行列交叉处解析出的证据，必须优先采用其中的值。
- 不得用其他 Table 的同名行、叙述性段落或相似数值覆盖该单元格证据。

【强制规则 5A：表格证据一致性】
- 如果参考资料中已经出现与问题直接对应的表格行、列和值，即使该片段没有重复显示完整的表号或题注，也应直接回答已确认的值。
- 不得在已经给出具体表格数值后，再说“资料未提供”“无法确认”或否认该数值属于用户指定的表格；如果只有部分子项有证据，只对缺失子项单独说明资料不足。

【其他要求】
- 表格以 Markdown 形式给出，数值问题请直接依据表格行列作答。
- 若问题涉及图片内容：如果参考片段没有标记为“图形坐标文字”的坐标证据，必须说明“该图内容未纳入文本检索范围”；如果存在这类坐标证据，只能依据其中明确出现的标签、数值和坐标范围作答，不能把它当作像素级图片识别。
- 若问题包含“多少、哪些、如何、管道、步骤”等多个事实维度，先在内部逐项核对问题要求，综合所有互补片段；不得因第一段已有概述就省略后续片段中的专有名词、工具名、阈值、数据规模、筛选条件或生成步骤。
- 若用户问题包含两个或以上事实维度，优先使用分点回答，并逐项覆盖参考资料中与问题直接相关的数字、阈值、工具/模型名称、实体和操作步骤；“流程概述”不能替代这些具体事实。
- 在回复“资料未提供相关信息”之前，必须逐一检查全部参考片段（包括后面编号的算法、附录和方法片段）；只要其中存在直接对应的初始化、残差、步长、循环类型、模型名或其他事实，就先回答该事实，不能因为前面的片段没有它而拒答整题。
- 复合问题中某一子项缺少证据时，回答其余有直接证据的子项，并明确指出仅缺少哪一项；不得把局部缺失扩大为整题拒答，也不得用常识补写未出现的细节。
- 若参考片段无法回答问题，请如实说明“资料未提供相关信息”，严禁编造。"""


def generate_mindmap(runtime: Runtime | None = None) -> str:
    runtime = runtime or get_runtime()
    if runtime.collection.count() == 0:
        return "📚 知识库为空，请先上传文档。"
    all_chunks = runtime.collection.get(include=["documents"])
    documents = all_chunks.get("documents") or []
    if not documents:
        return "无法读取文档内容。"
    context = "\n\n".join(documents[:15])
    try:
        response = runtime.client.chat.completions.create(
            model=runtime.config.deepseek_model,
            messages=[
                {"role": "system", "content": "你是一位顶级学术助教。请根据提供的课程资料，生成一份层级清晰、结构完整的学习大纲。"},
                {"role": "user", "content": f"请基于以下资料生成Markdown格式的层级大纲（使用 # ## ### - 表示层级），不要包含任何开场白或结尾总结，直接输出大纲结构。\n\n资料内容：\n{context}"},
            ],
            temperature=0.3,
            max_tokens=2000,
        )
        return response.choices[0].message.content
    except Exception as exc:
        return f"❌ 生成大纲失败：{exc}"


def generate_quiz(runtime: Runtime | None = None) -> str:
    runtime = runtime or get_runtime()
    if runtime.collection.count() == 0:
        return "📚 知识库为空，请先上传文档。"
    all_chunks = runtime.collection.get(include=["documents"])
    documents = all_chunks.get("documents") or []
    if not documents:
        return "无法读取文档内容。"
    sample_chunks = random.sample(documents, min(8, len(documents)))
    try:
        response = runtime.client.chat.completions.create(
            model=runtime.config.deepseek_model,
            messages=[
                {"role": "system", "content": "你是一个严谨的大学教师。请根据资料出5道单项选择题，用于考察学生对知识的掌握程度。"},
                {"role": "user", "content": "请根据以下资料，生成5道单项选择题。\n输出格式：第1题：[题目]\nA. [A] B. [B] C. [C] D. [D]\n答案：X\n解析：[解释]\n\n资料内容：\n" + "\n\n".join(sample_chunks)},
            ],
            temperature=0.4,
            max_tokens=2000,
        )
        return response.choices[0].message.content
    except Exception as exc:
        return f"❌ 出题失败：{exc}"


def build_demo(runtime: Runtime | None = None) -> Any:
    """Build the UI around an explicitly supplied runtime."""

    import gradio as gr

    runtime = runtime or get_runtime()
    with gr.Blocks(title="AI 大学生学习工作台", theme=gr.themes.Soft()) as demo:
        gr.Markdown("# 🎓 AI 大学生学习工作台")
        gr.Markdown("上传你的课件或论文，用 AI 帮你学！")
        with gr.Tab("📤 上传文档"):
            file_input = gr.File(label="选择文档", file_types=[".pdf", ".txt", ".docx"])
            upload_output = gr.Textbox(label="上传状态", lines=3)
            upload_button = gr.Button("添加到知识库")
            upload_button.click(lambda file: upload_file(file, runtime), inputs=file_input, outputs=upload_output)
            gr.Markdown(f"**当前知识库文本块数：** {runtime.collection.count()}")
        with gr.Tab("💬 智能问答"):
            gr.ChatInterface(
                fn=lambda message, history: query_knowledge(message, history, False, runtime),
                title="📖 基于文档的问答",
                description="输入问题，AI会从已上传的文档中检索答案。",
                chatbot=gr.Chatbot(height=450),
                textbox=gr.Textbox(placeholder="例如：这篇论文的核心创新点是什么？", scale=7),
            )
        with gr.Tab("🧠 生成学习大纲"):
            gr.Markdown("### 一键生成层级学习大纲（自动转为脑图结构）")
            button = gr.Button("🚀 生成大纲与脑图")
            output = gr.Markdown(label="📋 大纲内容", value="点击上方按钮生成...")
            button.click(lambda: generate_mindmap(runtime), inputs=[], outputs=output)
        with gr.Tab("📝 智能出题"):
            gr.Markdown("### 基于当前知识库自动生成练习题（含解析）")
            button = gr.Button("📝 生成5道选择题")
            output = gr.Markdown(label="📋 题目与解析", value="点击上方按钮生成...")
            button.click(lambda: generate_quiz(runtime), inputs=[], outputs=output)
    return demo


def chat_respond(message: str, history: Any) -> str:
    return query_knowledge(message, history, return_contexts=False)


if __name__ == "__main__":
    demo = build_demo(create_runtime())
    demo.launch()
